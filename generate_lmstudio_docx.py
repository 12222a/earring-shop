from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DESKTOP = Path(r"C:\Users\Administrator\Desktop")
OUTPUT = DESKTOP / "LM Studio Qwen3.5-9B 本地部署与使用指南.docx"


def set_cell_text(cell, text):
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(10.5)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.append(r_fonts)

    new_run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_code_block(document, text):
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["CodeBlock"]
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9.5)


def bullet(document, text):
    p = document.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)


def numbered(document, text):
    p = document.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)


def main():
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    if "CodeBlock" not in document.styles:
        style = document.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Consolas"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        style.font.size = Pt(9.5)
        style.paragraph_format.space_before = Pt(3)
        style.paragraph_format.space_after = Pt(3)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LM Studio 与 Qwen3.5-9B 本地部署与使用指南")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 73, 125)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = subtitle.add_run(
        f"生成日期：{date(2026, 3, 12).isoformat()}    适用机器：RTX 3060 12GB / 32GB RAM / Ryzen 5 5600"
    )
    srun.font.name = "Microsoft YaHei"
    srun._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    srun.font.size = Pt(10.5)

    document.add_paragraph(
        "本文档基于当前这台电脑已经完成的实际部署结果编写，目标是让你后续可以直接使用 LM Studio 本地模型，"
        "并让其他软件通过兼容 OpenAI 的方式调用它。"
    )

    document.add_heading("1. 当前机器与部署结果", level=1)
    env_table = document.add_table(rows=1, cols=2)
    env_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    env_table.style = "Table Grid"
    set_cell_text(env_table.rows[0].cells[0], "项目")
    set_cell_text(env_table.rows[0].cells[1], "当前值")
    env_rows = [
        ("操作系统", "Windows（当前工作目录：D:\\codex）"),
        ("LM Studio 版本", "0.4.6+1"),
        ("显卡", "NVIDIA GeForce RTX 3060 12GB"),
        ("内存", "32GB"),
        ("CPU", "AMD Ryzen 5 5600（6 核 12 线程）"),
        ("已部署模型", "Qwen3.5-9B-Q4_K_M.gguf"),
        ("视觉投影文件", "mmproj-Qwen3.5-9B-BF16.gguf"),
        ("模型主文件路径", r"C:\Users\Administrator\.lmstudio\models\lmstudio-community\Qwen3.5-9B-GGUF\Qwen3.5-9B-Q4_K_M.gguf"),
        ("视觉投影路径", r"C:\Users\Administrator\.lmstudio\models\lmstudio-community\Qwen3.5-9B-GGUF\mmproj-Qwen3.5-9B-BF16.gguf"),
        ("模型是否已被 LM Studio 索引", "是，已写入 model-index-cache.json"),
        ("文档生成时本地 API Server 状态", "未运行，需要在 LM Studio 中手动启动"),
    ]
    for left, right in env_rows:
        row = env_table.add_row().cells
        set_cell_text(row[0], left)
        set_cell_text(row[1], right)

    document.add_paragraph()
    bullet(document, "Gemma 3 4B 是 LM Studio 首次启动时自动提供的示例模型，可用，但不是本次主用模型。")
    bullet(document, "Qwen3.5-9B 已完成本地文件部署，LM Studio 内部索引已刷新。")
    bullet(document, "Qwen3.5-9B 是带视觉能力的模型，视觉功能依赖 mmproj 文件；该文件也已准备完成。")

    document.add_heading("2. 为什么选这个量化版本", level=1)
    document.add_paragraph(
        "本机最合适的版本是 Qwen3.5-9B 的 GGUF 量化文件 Q4_K_M。原因是："
    )
    bullet(document, "9B 体量足够做日常对话、代码辅助、文档理解和图文问答。")
    bullet(document, "Q4_K_M 在 12GB 显存上比较平衡，质量明显好于小模型，同时不会像更高量化精度那样容易顶满显存。")
    bullet(document, "与 RTX 3060 12GB 配合时，Q4_K_M 比 Q6_K / Q8_0 更稳，更适合长期本地主力使用。")

    document.add_heading("3. 在 LM Studio 中第一次加载模型", level=1)
    numbered(document, "打开 LM Studio。")
    numbered(document, "进入模型选择界面，选择 Qwen3.5 9B。")
    numbered(document, "按照下表设置加载参数。")
    numbered(document, "点击“加载模型”。")
    numbered(document, "加载成功后，新建一个 Chat，输入一条测试消息，例如“你好，请用三句话介绍自己”。")

    param_table = document.add_table(rows=1, cols=3)
    param_table.style = "Table Grid"
    param_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(param_table.rows[0].cells[0], "参数")
    set_cell_text(param_table.rows[0].cells[1], "建议值")
    set_cell_text(param_table.rows[0].cells[2], "原因")
    param_rows = [
        ("上下文长度", "4096（稳定后可升到 8192）", "第一次加载最稳；上下文越长，KV Cache 占用越高。"),
        ("GPU 卸载", "32", "Qwen3.5-9B 元数据中为 32 层，拉满表示尽量全放 GPU，速度最好。"),
        ("CPU 线程池大小", "6", "你的 CPU 是 6 核，通常用物理核心数更稳。"),
        ("评估批处理大小", "512", "对 12GB 显存比较稳；更大可能更快，但更容易吃满显存。"),
        ("Max Concurrent Predictions", "1", "本地单用户使用时最合适；4 会额外吃显存和 KV Cache。"),
        ("Unified KV Cache", "开", "默认建议开启。"),
        ("将 KV 缓存卸载到 GPU 内存", "开", "有足够显存时速度更好。"),
        ("保持模型在内存中", "开", "避免每次切换聊天都重新加载模型。"),
        ("尝试 mmap()", "开", "模型文件映射加载通常更省内存、更快。"),
        ("快速注意力", "开", "通常更快。"),
        ("RoPE 频率基 / 比例", "自动", "除非有特殊长上下文实验需求，否则保持默认。"),
        ("K/V 缓存量化类型", "先不要改", "默认最稳；后续如果追求更长上下文再考虑调。"),
    ]
    for item, value, why in param_rows:
        row = param_table.add_row().cells
        set_cell_text(row[0], item)
        set_cell_text(row[1], value)
        set_cell_text(row[2], why)

    document.add_paragraph()
    bullet(document, "如果第一次加载报显存不足，优先把“上下文长度”降到 4096，必要时把“评估批处理大小”降到 256。")
    bullet(document, "如果模型稳定运行，可以将上下文长度从 4096 提到 8192。")

    document.add_heading("4. 在 LM Studio 内如何实际使用", level=1)
    bullet(document, "普通聊天：直接在 Chat 页输入问题即可。")
    bullet(document, "文档总结：把文本或文件内容贴入聊天窗口，让模型做摘要、问答或改写。")
    bullet(document, "写代码：把需求、错误信息或代码片段贴进去，请它解释、修复或生成。")
    bullet(document, "图文问答：在支持图片输入的聊天窗口里上传图片，然后让模型分析图片内容。")
    bullet(document, "多轮对话：保持同一个 Chat，可以持续追问上下文。")

    document.add_heading("5. 让其他软件调用本地模型", level=1)
    document.add_paragraph(
        "最常见的方式，是把 LM Studio 当成一个本地版 OpenAI API 服务。只要别的软件支持 OpenAI 或 OpenAI-compatible API，通常都能接。"
    )
    numbered(document, "在 LM Studio 中打开本地 API Server（Developer / Local Server 页面）。")
    numbered(document, "确认监听地址与端口。官方文档默认是 localhost:1234。")
    numbered(document, "如果软件和 LM Studio 在同一台电脑上，Base URL 填 http://localhost:1234/v1。")
    numbered(document, "如果你要让局域网其他设备调用，需要打开 Serve on Local Network，并改用 http://你的电脑IP:1234/v1。")
    numbered(document, "如果你开启了认证，就把 LM Studio 生成的 token 作为 API Key；如果没开认证，多数客户端可以随便填一个占位值，例如 lm-studio。")

    api_table = document.add_table(rows=1, cols=2)
    api_table.style = "Table Grid"
    api_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(api_table.rows[0].cells[0], "字段")
    set_cell_text(api_table.rows[0].cells[1], "建议填写")
    api_rows = [
        ("Provider / 供应商", "OpenAI-compatible / OpenAI 兼容"),
        ("Base URL", "http://localhost:1234/v1"),
        ("API Key", "lm-studio（如果未启用认证）"),
        ("推荐模型名", "先试 qwen3.5-9b；若软件要求完整 id，就用 /v1/models 返回值"),
        ("备选模型标识", "lmstudio-community/qwen3.5-9b"),
        ("是否需要先手工加载", "如果启用了 Just-in-Time Model Loading，则可自动加载；否则先在 LM Studio 中手工加载"),
    ]
    for k, v in api_rows:
        row = api_table.add_row().cells
        set_cell_text(row[0], k)
        set_cell_text(row[1], v)

    document.add_paragraph()
    document.add_paragraph("常用接口：")
    bullet(document, "GET /v1/models：列出可用模型")
    bullet(document, "POST /v1/chat/completions：兼容 OpenAI Chat Completions")
    bullet(document, "POST /v1/responses：兼容 OpenAI Responses API")
    bullet(document, "POST /v1/embeddings：如果后续使用本地嵌入模型")

    document.add_paragraph("示例 1：先列出模型")
    add_code_block(
        document,
        "curl http://localhost:1234/v1/models",
    )

    document.add_paragraph("示例 2：使用 curl 做一次聊天请求")
    add_code_block(
        document,
        'curl http://localhost:1234/v1/chat/completions ^\n'
        '  -H "Content-Type: application/json" ^\n'
        '  -H "Authorization: Bearer lm-studio" ^\n'
        '  -d "{\\"model\\":\\"qwen3.5-9b\\",\\"messages\\":[{\\"role\\":\\"user\\",\\"content\\":\\"你好，请用三句话介绍自己\\"}]}"',
    )

    document.add_paragraph("示例 3：使用 Python（OpenAI SDK）调用")
    add_code_block(
        document,
        "from openai import OpenAI\n\n"
        'client = OpenAI(base_url="http://localhost:1234/v1", api_key="lm-studio")\n\n'
        "print(client.models.list())\n\n"
        "resp = client.chat.completions.create(\n"
        '    model="qwen3.5-9b",\n'
        '    messages=[{"role": "user", "content": "你好"}],\n'
        ")\n\n"
        "print(resp.choices[0].message.content)\n",
    )

    document.add_heading("6. 其他常见软件如何接", level=1)
    bullet(document, "Cherry Studio：选择 OpenAI 兼容供应商，填 Base URL 和 API Key，再选择 qwen3.5-9b。")
    bullet(document, "Open WebUI / AnythingLLM：把 LM Studio 当成 OpenAI 接口源，Base URL 填 http://localhost:1234/v1。")
    bullet(document, "Cline / Roo Code / 其他代码插件：只要支持 OpenAI 兼容接口，通常都可以接本地 LM Studio。")
    bullet(document, "如果接入软件里看不到模型，先调用 /v1/models 看实际返回的 model id，再把那个值填进软件。")

    document.add_heading("7. Just-in-Time Model Loading 的理解", level=1)
    document.add_paragraph(
        "LM Studio 官方文档提到，如果启用了 Just-in-Time Model Loading（JIT 模型按需加载），"
        "那么 API 请求可以直接指定一个已经下载到本地但当前未手工加载的模型，LM Studio 会在收到请求时自动加载它。"
        "这对“让别的软件直接调本地模型”非常有用。"
    )
    bullet(document, "优点：外部软件不需要你每次先手工点“加载模型”。")
    bullet(document, "前提：模型已经下载到本地，并且 LM Studio 的本地服务已经启动。")
    bullet(document, "如果你不想自动加载，也可以手工先在 LM Studio 里把 Qwen3.5-9B 加载进内存。")

    document.add_heading("8. 故障排查", level=1)
    bullet(document, "模型能在 LM Studio 里聊天，但别的软件连不上：通常是本地 API Server 没启动。先去 Developer / Local Server 页面启动它。")
    bullet(document, "软件提示 404 或 model not found：先请求 http://localhost:1234/v1/models，确认真实的模型 id。")
    bullet(document, "模型加载时显存不足：优先降低上下文长度，其次把评估批处理大小从 512 降到 256。")
    bullet(document, "图像输入不工作：确认 mmproj 文件存在且使用的是 Qwen3.5 9B，而不是示例模型 Gemma。")
    bullet(document, "局域网其他设备调用失败：确认已打开 Serve on Local Network，并检查 Windows 防火墙。")
    bullet(document, "本机当前状态说明：文档生成时 lms server status 仍显示 The server is not running，因此 API 使用部分需要你在界面中手工启动服务后再调用。")

    document.add_heading("9. 本机关键文件与路径", level=1)
    path_table = document.add_table(rows=1, cols=2)
    path_table.style = "Table Grid"
    path_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(path_table.rows[0].cells[0], "用途")
    set_cell_text(path_table.rows[0].cells[1], "路径")
    path_rows = [
        ("LM Studio 程序目录", r"C:\Users\Administrator\AppData\Local\Programs\LM Studio"),
        ("LM Studio 主设置", r"C:\Users\Administrator\.lmstudio\settings.json"),
        ("模型索引缓存", r"C:\Users\Administrator\.lmstudio\.internal\model-index-cache.json"),
        ("Qwen3.5 主模型", r"C:\Users\Administrator\.lmstudio\models\lmstudio-community\Qwen3.5-9B-GGUF\Qwen3.5-9B-Q4_K_M.gguf"),
        ("Qwen3.5 视觉投影", r"C:\Users\Administrator\.lmstudio\models\lmstudio-community\Qwen3.5-9B-GGUF\mmproj-Qwen3.5-9B-BF16.gguf"),
        ("LM Studio 日志目录", r"C:\Users\Administrator\AppData\Roaming\LM Studio\logs"),
    ]
    for left, right in path_rows:
        row = path_table.add_row().cells
        set_cell_text(row[0], left)
        set_cell_text(row[1], right)

    document.add_heading("10. 官方参考链接", level=1)
    p = document.add_paragraph()
    add_hyperlink(p, "LM Studio OpenAI-compatible API", "https://lmstudio.ai/docs/developer/openai-compat")
    p = document.add_paragraph()
    add_hyperlink(p, "LM Studio Server Settings", "https://lmstudio.ai/docs/developer/core/server/settings")
    p = document.add_paragraph()
    add_hyperlink(p, "LM Studio Headless / JIT Loading", "https://lmstudio.ai/docs/developer/core/headless")
    p = document.add_paragraph()
    add_hyperlink(p, "LM Studio REST API", "https://lmstudio.ai/docs/developer/rest")

    document.add_paragraph()
    end_note = document.add_paragraph()
    end_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_run = end_note.add_run("本文档已根据本机实际部署结果生成，可直接作为后续使用说明。")
    end_run.font.name = "Microsoft YaHei"
    end_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    end_run.font.bold = True

    document.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
