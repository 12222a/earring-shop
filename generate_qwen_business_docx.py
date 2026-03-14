from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DESKTOP = Path(r"C:\Users\Administrator\Desktop")
OUTPUT = DESKTOP / "千问（Qwen）业务应用场景研究与落地建议.docx"


def style_run(run, size=10.5, bold=False, color=None, font_name="Microsoft YaHei"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_text(cell, text):
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            style_run(run)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.append(r_fonts)
    new_run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_code_block(document, text):
    p = document.add_paragraph(style="CodeBlock")
    run = p.add_run(text)
    style_run(run, size=9.5, font_name="Consolas")


def bullet(document, text):
    p = document.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    style_run(run)


def numbered(document, text):
    p = document.add_paragraph(style="List Number")
    run = p.add_run(text)
    style_run(run)


def paragraph(document, text, bold_prefix=None):
    p = document.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        run1 = p.add_run(bold_prefix)
        style_run(run1, bold=True)
        run2 = p.add_run(text[len(bold_prefix):])
        style_run(run2)
    else:
        run = p.add_run(text)
        style_run(run)
    return p


def main():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    if "CodeBlock" not in doc.styles:
        style = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Consolas"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        style.font.size = Pt(9.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("千问（Qwen）业务应用场景研究与落地建议")
    style_run(r, size=20, bold=True, color=(31, 73, 125))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("基于官方资料整理 | 搜索日期：2026-03-12")
    style_run(r, size=10.5)

    paragraph(
        doc,
        "说明：你前面提到的“前问”，这里按“千问 / Qwen 家族”理解。本文尽量只采用官方或一手资料；"
        "其中“业务上可以怎么用”的部分，有些是官方明确给出的产品场景，有些是基于官方能力说明做出的业务推断，文中会明确标注。",
    )

    doc.add_heading("一、结论先看", level=1)
    bullet(doc, "如果从业务落地角度看，千问最适合优先切入的方向，不是“做一个万能聊天机器人”，而是做 8 类高价值场景：智能客服、企业知识库问答、文档/单据解析、会议与录音洞察、销售质检、研发提效、金融垂直助手、内容生产。")
    bullet(doc, "如果你已经部署的是本地 Qwen3.5-9B，那么它最适合做部门级、单机级、低并发的内部应用和 PoC，不太适合直接承担高并发生产流量。")
    bullet(doc, "从官方产品布局看，千问并不是只有一个“文本大模型”，而是已经形成了文本、代码、多模态、Embedding/Rerank、安全护栏、金融垂直、音视频洞察等组合能力。")

    doc.add_heading("二、官方资料里能明确看到的能力版图", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], "能力方向")
    set_cell_text(table.rows[0].cells[1], "官方依据")
    set_cell_text(table.rows[0].cells[2], "能支撑的业务动作")
    set_cell_text(table.rows[0].cells[3], "是否更适合本地 9B")

    rows = [
        ("通用文本推理与多语言", "Qwen3 支持 119 种语言与方言，并强化了 agent 能力", "问答、写作、翻译、流程助手、跨语种支持", "适合"),
        ("代码与研发 Agent", "Qwen3-Coder 面向 agentic coding 和真实软件工程任务", "代码生成、修复、测试、Repo 问答、研发助手", "适合轻中度场景"),
        ("视觉与文档理解", "Qwen2.5-VL 支持 OCR、票据表单结构化输出、视觉定位、长视频理解", "发票/表单抽取、资质审核、质检、图文问答", "适合轻量单据与图文问答"),
        ("Embedding 与 Rerank", "Qwen3 Embedding 系列面向文本嵌入、检索、重排", "企业知识库、语义搜索、RAG 召回", "适合作为本地知识库底座的一部分"),
        ("安全护栏", "Qwen3Guard 支持输入与输出的安全分类和实时检测", "内容审核、敏感指令拦截、审计", "更适合服务端流程化部署"),
        ("音视频理解", "通义听悟可做转写、摘要、待办、问答、客户服务/会议分析", "会议纪要、客服录音分析、销售复盘", "本地 9B 不直接覆盖音频主流程"),
        ("金融垂直", "通义点金定位为基于千问的金融领域模型应用", "金融智能体、报告助手、文档解析", "建议用云端或专用产品"),
        ("图像生成与编辑", "Qwen-Image / Qwen-Image-Edit 强调中文文本渲染和精细编辑", "海报、PPT、广告素材、图像改版", "不建议用你当前本地 9B 承担"),
    ]
    for row_data in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(row[idx], value)

    doc.add_heading("三、千问具体能落在哪些业务上", level=1)

    doc.add_heading("1. 智能客服与自助服务", level=2)
    paragraph(doc, "这是千问最明确、最成熟的业务方向之一。阿里云官方的“智能对话机器人（千问版）”直接说明，它是基于千问训练的客服领域大模型，可部署在网站、APP、智能硬件等前端入口，用于提升服务体验与效率。")
    bullet(doc, "可做什么：FAQ 自助问答、售前咨询、售后分流、订单/流程解释、政策说明、网站资料问答。")
    bullet(doc, "业务价值：减少人工坐席压力，提升 7x24 响应能力。")
    bullet(doc, "适配形态：官网机器人、App 内客服助手、企业微信/钉钉客服入口。")
    bullet(doc, "对你当前本地 Qwen3.5-9B 的判断：适合先做内部客服知识助手或低并发外部演示版；正式大规模客服更建议接云端模型或更完整的产品形态。")
    paragraph(doc, "来源：阿里云官方“智能对话机器人（千问版）上线公告”。")

    doc.add_heading("2. 企业知识库问答与内部专家助手", level=2)
    paragraph(doc, "这类场景不是简单聊天，而是“用企业自己的资料回答问题”。官方资料里，阿里云已有“Qwen3 + 联网搜索 + RAG”的应用模板；Qwen3 Embedding 系列则专门面向文本嵌入、检索和重排。")
    bullet(doc, "可做什么：制度问答、SOP 查询、产品手册问答、研发文档检索、销售资料助手、培训知识助手。")
    bullet(doc, "业务价值：把“找资料”变成“问问题”，减少跨部门沟通成本。")
    bullet(doc, "推荐组合：Qwen3/Qwen3.5 做回答，Qwen3 Embedding + Reranker 做召回和重排。")
    bullet(doc, "对你当前本地 Qwen3.5-9B 的判断：这是最适合你现在就落地的方向之一。")
    paragraph(doc, "来源：Qwen3 Embedding 官方博客；PAI 官方“LangStudio x Qwen3-RAG 联网搜索 AI 智能问答应用”。")

    doc.add_heading("3. 文档理解、表单抽取、资质审核、财务单据解析", level=2)
    paragraph(doc, "Qwen2.5-VL 官方明确写到，它支持 OCR、文本定位、票据/表单/表格等内容的结构化输出，并点名适用于 qualification review 和 financial business。")
    bullet(doc, "可做什么：发票/报销单抽取、合同关键字段提取、资质文件初审、表单自动录入、图文材料归档。")
    bullet(doc, "业务价值：替代大量人工录入与人工初筛。")
    bullet(doc, "适配行业：财务、法务、人事、政务、教育、保险、供应链。")
    bullet(doc, "对你当前本地 Qwen3.5-9B 的判断：适合轻量图文问答和少量单据实验；若要做批量 OCR、复杂票据处理、稳定结构化抽取，建议走更强的视觉模型或云端服务。")
    paragraph(doc, "来源：Qwen2.5-VL 官方博客。")
    paragraph(doc, "说明：这里把模型能力映射到“资质审核、单据处理、合同抽取”等业务流程，属于基于官方能力的业务推断。")

    doc.add_heading("4. 会议纪要、访谈总结、客服录音分析、销售复盘", level=2)
    paragraph(doc, "阿里云官方“通义听悟”直接给出了适用场景：会议、拜访、面试、客户服务、访谈、路演等。它能够在转写、翻译、说话人分离的基础上，生成摘要、章节速览、发言总结、待办事项、问答、关键词。")
    bullet(doc, "可做什么：会议纪要自动生成、客户拜访总结、客服录音提炼、面试记录结构化、销售复盘。")
    bullet(doc, "业务价值：让音视频内容可检索、可复盘、可被管理层消费。")
    bullet(doc, "典型落地：销售团队周复盘、客服中心质检、HR 面试归档、投研访谈整理。")
    paragraph(doc, "来源：阿里云官方“工作学习 AI 助手通义听悟”。")

    doc.add_heading("5. 销售质检、门店话术分析、服务流程洞察", level=2)
    paragraph(doc, "阿里云官方还给了一个更垂直的实例：汽车销售服务洞察 Agent。文档中说明，默认就有 11 大类 74 项洞察项，涵盖开场介绍、意向确认、现场接待、需求沟通等，还可接音频、视频和文本内容。")
    bullet(doc, "可做什么：销售对话质检、关键话术命中分析、服务流程执行度检查、培训反馈。")
    bullet(doc, "业务价值：把“听录音”变成“读洞察结果”，适合门店管理、呼叫中心、销售培训。")
    bullet(doc, "可迁移行业：汽车只是官方例子，同样思路可迁移到保险、电销、教培顾问、医美咨询、房地产销售。")
    paragraph(doc, "来源：阿里云官方“通义听悟 Agent - 汽车销售服务洞察”使用指南。")
    paragraph(doc, "说明：跨行业迁移这一点是基于官方汽车场景做出的业务推断。")

    doc.add_heading("6. 软件研发提效与 AI 开发助手", level=2)
    paragraph(doc, "Qwen3-Coder 官方明确面向 agentic coding，且强调真实软件工程任务、工具使用、多轮规划、测试反馈和 Repo 级长上下文。")
    bullet(doc, "可做什么：代码生成、单元测试、Bug 修复、脚本编写、接口文档生成、PR 总结、代码解释、项目问答。")
    bullet(doc, "业务价值：缩短研发交付时间，让中小团队更像“有了一个 24 小时技术搭子”。")
    bullet(doc, "适用团队：开发、测试、数据工程、运维、实施顾问。")
    bullet(doc, "对你当前本地 Qwen3.5-9B 的判断：很适合做个人研发助手、命令行辅助、脚本生成和小仓库问答；若要做复杂仓库级 Agent，Qwen3-Coder 或更强云端版本更合适。")
    paragraph(doc, "来源：Qwen3-Coder 官方博客。")

    doc.add_heading("7. 金融业务智能体与金融文档助手", level=2)
    paragraph(doc, "阿里云官方“通义点金”直接定位为基于千问大模型、专注于金融领域的模型应用，可通过金融业务智能体和金融模型提供服务。")
    bullet(doc, "可做什么：研报摘要、投顾问答、合规文本整理、金融文档解析、业务坐席辅助。")
    bullet(doc, "业务价值：减少金融从业者在信息整理和文本分析上的重复劳动。")
    bullet(doc, "落地建议：金融业务优先用官方金融垂直产品或云端受控环境，不建议仅靠本地 9B 承担关键生产职责。")
    paragraph(doc, "来源：阿里云官方“通义点金”产品简介。")

    doc.add_heading("8. 内容营销、海报、PPT 与视觉物料生产", level=2)
    paragraph(doc, "Qwen-Image 官方强调中文和英文复杂文本渲染能力，并展示了海报和 PPT 页面生成；Qwen-Image-Edit 则强调双语文本编辑和局部精准修改。")
    bullet(doc, "可做什么：海报、活动页、商品图文物料、PPT 页、宣传图改版、图片中的中文文案修正。")
    bullet(doc, "业务价值：降低平面设计和内容制作的门槛，缩短营销物料出稿时间。")
    bullet(doc, "适用团队：市场、品牌、运营、电商、培训。")
    paragraph(doc, "来源：Qwen-Image、Qwen-Image-Edit 官方博客。")

    doc.add_heading("9. 安全审核、风险拦截、内容合规", level=2)
    paragraph(doc, "Qwen3Guard 官方说明它能对用户输入 prompt 和模型响应进行安全分类，且支持实时流式检测。")
    bullet(doc, "可做什么：越狱提示拦截、敏感话题识别、不当回复检测、模型输出审计。")
    bullet(doc, "业务价值：让企业在上线 AI 助手时多一层安全防护。")
    bullet(doc, "典型位置：放在客服机器人前后、企业知识库问答前后、内容生成流水线里。")
    paragraph(doc, "来源：Qwen3Guard 官方博客。")

    doc.add_heading("10. 私有化、敏感数据、保密型 AI 应用", level=2)
    paragraph(doc, "阿里云官方的机密计算方案里，直接给出 Qwen 在 finance and healthcare、private enterprise knowledge bases、commercial model protection 这些场景的适配方式。")
    bullet(doc, "可做什么：处理医疗、金融、研发、法务等敏感数据时，提供更强的数据保护。")
    bullet(doc, "业务价值：适合对合规和保密要求极高的企业。")
    bullet(doc, "你可以怎么理解：不是“千问只能上公有云”，而是官方已经考虑到了保密部署场景。")
    paragraph(doc, "来源：Alibaba Cloud Linux 官方文档“Deploy a private Qwen-7B-Chat model protected by Intel TDX using the Confidential AI solution”。")

    doc.add_heading("四、哪些业务最适合你现在这台机器先做", level=1)
    paragraph(doc, "以下判断不是来自官方文档，而是结合你当前已部署的本地 Qwen3.5-9B、RTX 3060 12GB 和 32GB 内存做出的落地建议。")
    good_table = doc.add_table(rows=1, cols=3)
    good_table.style = "Table Grid"
    good_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(good_table.rows[0].cells[0], "优先级")
    set_cell_text(good_table.rows[0].cells[1], "最适合先做的业务")
    set_cell_text(good_table.rows[0].cells[2], "原因")
    good_rows = [
        ("P1", "企业知识库问答 / 内部专家助手", "文本为主、低并发、最容易体现价值。"),
        ("P1", "文档总结 / 制度问答 / SOP 助手", "部署简单，几乎不需要复杂多模态链路。"),
        ("P1", "个人或小团队研发助手", "你已经本地部署完成，马上能用。"),
        ("P2", "轻量图文问答 / 单据抽取 PoC", "Qwen3.5 具备视觉能力，但更适合验证方案而非高吞吐生产。"),
        ("P2", "营销文案生成 / 邮件 / 报告草稿", "文本生成成本低，ROI 清晰。"),
        ("P3", "正式客服机器人外放", "能做演示版，但生产环境还要补安全、监控、并发与知识治理。"),
        ("P3", "大规模录音分析 / 视频洞察", "本地 9B 不是最佳主力，建议走专门音视频产品链路。"),
    ]
    for item in good_rows:
        row = good_table.add_row().cells
        for idx, value in enumerate(item):
            set_cell_text(row[idx], value)

    doc.add_heading("五、如果你要从业务角度推进，我建议的落地顺序", level=1)
    numbered(doc, "先做一个内部知识库助手。目标不是“像 ChatGPT 一样聊天”，而是让它回答你公司最常见、最重复、最耗时间的问题。")
    numbered(doc, "再做一个文档处理或 SOP 助手。比如招投标材料提要、制度问答、合同关键字段提取。")
    numbered(doc, "如果你团队有研发人员，再把它做成代码助手。这样最容易直接看到效率提升。")
    numbered(doc, "如果你所在业务强依赖录音、会议、客服或销售对话，再引入听悟类能力做“总结与洞察”。")
    numbered(doc, "最后才考虑大规模客户外放、复杂合规、跨系统自动化。因为这些环节需要更多工程能力，不只是模型本身。")

    doc.add_heading("六、一个可执行的业务判断框架", level=1)
    bullet(doc, "如果你的业务本质是“找资料、读文档、解释规则”，优先考虑 Qwen + RAG。")
    bullet(doc, "如果你的业务本质是“看图片、看表单、看票据”，优先考虑 VL/多模态路线。")
    bullet(doc, "如果你的业务本质是“听录音、开会、拜访复盘”，优先考虑听悟 / 音视频分析路线。")
    bullet(doc, "如果你的业务本质是“写代码、修脚本、查仓库”，优先考虑 Qwen3-Coder。")
    bullet(doc, "如果你的业务本质是“高风险内容生成”，一定补上 Qwen3Guard 或等效安全层。")
    bullet(doc, "如果你的业务数据极其敏感，优先规划私有化或机密计算形态。")

    doc.add_heading("七、对“能不能商用”的一句提醒", level=1)
    paragraph(
        doc,
        "从官方资料看，千问已经不是实验型模型，而是有完整产品化路线和行业化组件的模型家族。"
        "但“能商用”不等于“拿一个模型就能上线”。真正的商用还要补齐知识库治理、权限、日志、审计、安全护栏、流程编排、监控和成本控制。",
    )

    doc.add_heading("八、本文使用的主要来源", level=1)
    sources = [
        ("Qwen3 官方博客", "https://qwenlm.github.io/blog/qwen3/"),
        ("Qwen3-Coder 官方博客", "https://qwenlm.github.io/blog/qwen3-coder/"),
        ("Qwen2.5-VL 官方博客", "https://qwenlm.github.io/blog/qwen2.5-vl/"),
        ("Qwen3 Embedding 官方博客", "https://qwenlm.github.io/blog/qwen3-embedding/"),
        ("Qwen3Guard 官方博客", "https://qwenlm.github.io/blog/qwen3guard/"),
        ("Qwen-Image 官方博客", "https://qwenlm.github.io/blog/qwen-image/"),
        ("智能对话机器人（千问版）上线公告", "https://help.aliyun.com/zh/beebot/intelligent-dialogue-robot-tongyi-edition-online-announcement"),
        ("通义听悟官方概览", "https://help.aliyun.com/zh/tingwu/"),
        ("通义听悟 Agent - 汽车销售服务洞察", "https://help.aliyun.com/zh/model-studio/tingwu-automotive-service-insights-guidelines"),
        ("通义点金官方简介", "https://help.aliyun.com/zh/model-studio/tongyi-dianjin-overview"),
        ("LangStudio x Qwen3-RAG 联网搜索应用", "https://help.aliyun.com/zh/pai/use-cases/langstudio-x-qwen3-rag-network-search-ai-intelligent-question-and-answer-application"),
        ("Confidential AI with Qwen 官方文档", "https://www.alibabacloud.com/help/en/alinux/deploy-privacy-qwen-7b-chat-model-protected-by-intel-tdx-with-confidential-ai-solution"),
    ]
    for name, url in sources:
        p = doc.add_paragraph()
        add_hyperlink(p, name, url)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("如果你下一步想做行业版，我建议继续细化成“你所在行业的千问业务地图”。")
    style_run(r, bold=True, color=(31, 73, 125))

    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
